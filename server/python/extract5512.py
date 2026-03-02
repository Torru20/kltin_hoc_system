#!/usr/bin/env python3
import sys, io, json, re
from docx import Document

# Đảm bảo in tiếng Việt trên Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def clean(s):
    return re.sub(r'\s+', ' ', s.strip())

def iter_all_paragraphs(doc):
    """Đọc toàn bộ đoạn văn (bao gồm cả trong bảng)."""
    texts = []
    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        texts.append(p.text.strip())
    return texts


def extract_khbd(path):
    doc = Document(path)
    paras = [clean(p) for p in iter_all_paragraphs(doc)]

    data = {
        "thong_tin": {"text": ""},
        "muc_tieu": {
            "kien_thuc": [],
            "nang_luc": {"dac_thu": [], "chung": []},
            "pham_chat": []
        },
        "thiet_bi": {"giao_vien": [], "hoc_sinh": []}
    }

    current_section = "thong_tin"
    current_sub = None
    thiet_bi_sub = None

    for t in paras:
        low = t.lower()

        # ======= 1️⃣ MỤC TIÊU =======
        if re.search(r'\bmục\s*tiêu\b', low):
            current_section = "muc_tieu"
            current_sub = None
            continue

        if current_section == "muc_tieu":
            # -- dò ranh giới các phần nhỏ
            if re.search(r'về\s*kiến\s*thức', low):
                current_sub = "kien_thuc"
                continue
            elif re.search(r'về\s*năng\s*lực', low):
                current_sub = "nang_luc"
                continue
            elif re.search(r'năng\s*lực\s*đặc\s*thù', low):
                current_sub = "nang_luc_dac_thu"
                continue
            elif re.search(r'năng\s*lực\s*chung', low):
                current_sub = "nang_luc_chung"
                continue
            elif re.search(r'về\s*phẩm\s*chất', low):
                current_sub = "pham_chat"
                continue
            elif re.search(r'thiết\s*bị\s*dạy\s*học', low):
                current_section = "thiet_bi"
                current_sub = None
                continue

            # --- ghi đúng vùng ---
            if current_sub == "kien_thuc":
                data["muc_tieu"]["kien_thuc"].append(t)
            elif current_sub == "nang_luc_dac_thu":
                data["muc_tieu"]["nang_luc"]["dac_thu"].append(t)
            elif current_sub == "nang_luc_chung":
                data["muc_tieu"]["nang_luc"]["chung"].append(t)
            elif current_sub == "pham_chat":
                data["muc_tieu"]["pham_chat"].append(t)
            continue


        # ======= 2️⃣ THIẾT BỊ DẠY HỌC =======
        if re.search(r'thiết\s*bị\s*dạy\s*học', low):
            current_section = "thiet_bi"
            thiet_bi_sub = None
            continue

        if current_section == "thiet_bi":
            # Dừng lại khi gặp Tiến trình dạy học
            if re.search(r'tiến\s*trình\s*dạy\s*học', low):
                current_section = None
                break

            # xác định phần con
            if re.search(r'giáo\s*viên', low):
                thiet_bi_sub = "gv"
                continue
            elif re.search(r'học\s*sinh', low):
                thiet_bi_sub = "hs"
                continue

            # ghi nội dung
            if thiet_bi_sub == "gv":
                data["thiet_bi"]["giao_vien"].append(t)
            elif thiet_bi_sub == "hs":
                data["thiet_bi"]["hoc_sinh"].append(t)
            else:
                # nếu chưa xác định thì mặc định cho giáo viên
                data["thiet_bi"]["giao_vien"].append(t)
            continue


        # ======= THÔNG TIN =======
        if current_section == "thong_tin":
            data["thong_tin"]["text"] += " " + t


    # --- Làm sạch ---
    data["thong_tin"]["text"] = clean(data["thong_tin"]["text"])
    for k in ["kien_thuc", "pham_chat"]:
        data["muc_tieu"][k] = [clean(x) for x in data["muc_tieu"][k]]
    data["muc_tieu"]["nang_luc"]["dac_thu"] = [clean(x) for x in data["muc_tieu"]["nang_luc"]["dac_thu"]]
    data["muc_tieu"]["nang_luc"]["chung"] = [clean(x) for x in data["muc_tieu"]["nang_luc"]["chung"]]
    for k in ["giao_vien", "hoc_sinh"]:
        data["thiet_bi"][k] = [clean(x) for x in data["thiet_bi"][k]]

    print(json.dumps(data, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    extract_khbd(sys.argv[1])
